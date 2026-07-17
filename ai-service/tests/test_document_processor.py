"""文档处理分页提取单测（M4-4 增强：前端预览真实翻页）。

验证 document_processor.extract_pages 对四种可上传格式（pdf/docx/txt/md）
返回真实/估算分页，且拼接回全文与原文一致、page_no 从 1 递增。

注：项目未引入 pytest-asyncio，async 协程用标准 asyncio.run 驱动（见 test_llm.py）。
"""
import asyncio
import os
import tempfile

import pytest

from services.document_processor import document_processor, CHARS_PER_PAGE, PageSegment, _clean_text
from services import document_processor as dp_module


def _write_tmp(suffix: str, content: str) -> str:
    # 二进制写入，忠实保留原文换行符（Windows 文本模式会把 \n 转成 \r\n，导致拼接回不一致）
    fd, path = tempfile.mkstemp(suffix=suffix)
    with os.fdopen(fd, "wb") as f:
        f.write(content.encode("utf-8"))
    return path


def test_extract_pages_txt_estimated():
    """TXT 按 CHARS_PER_PAGE 估算分页，拼接回全文一致。"""
    text = "知识库正文内容。" * 200
    path = _write_tmp(".txt", text)
    try:
        pages = asyncio.run(document_processor.extract_pages(path, "txt"))
        assert len(pages) >= 1
        assert "".join(p.text for p in pages) == text
        assert [p.page for p in pages] == list(range(1, len(pages) + 1))
        assert len(pages) == max(1, (len(text) + CHARS_PER_PAGE - 1) // CHARS_PER_PAGE)
    finally:
        os.unlink(path)


def test_extract_pages_md_estimated():
    """MD 同 TXT 走估算分页。"""
    text = "# 标题\n" + "段落内容。" * 150
    path = _write_tmp(".md", text)
    try:
        pages = asyncio.run(document_processor.extract_pages(path, "md"))
        assert "".join(p.text for p in pages) == text
        assert [p.page for p in pages] == list(range(1, len(pages) + 1))
    finally:
        os.unlink(path)


def test_extract_pages_docx_estimated():
    """DOCX 提取正文 + 表格，估算分页，page_no 递增。"""
    from docx import Document

    path = _write_tmp(".docx", "")
    try:
        doc = Document()
        doc.add_paragraph("第一段内容。" * 50)
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "单元格A"
        table.cell(0, 1).text = "单元格B"
        doc.save(path)
        pages = asyncio.run(document_processor.extract_pages(path, "docx"))
        joined = "".join(p.text for p in pages)
        assert "第一段内容" in joined
        assert "单元格A" in joined and "单元格B" in joined
        assert [p.page for p in pages] == list(range(1, len(pages) + 1))
    finally:
        os.unlink(path)


def test_extract_pages_missing_file():
    """文件不存在应抛 FileNotFoundError，由路由捕获并降级。"""
    with pytest.raises(FileNotFoundError):
        asyncio.run(document_processor.extract_pages("/nonexistent/x.txt", "txt"))


def test_extract_pdf_falls_back_to_pdfplumber(monkeypatch):
    """PyMuPDF 不可用（如 _mupdf.pyd 因 DLL 加载失败）时降级到纯 Python 的 pdfplumber。

    通过强制 ``import fitz`` 抛 ImportError 并注入假 pdfplumber 模块，确定性验证降级分支。
    """
    import builtins
    import sys
    import types

    real_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name == "fitz" or name.startswith("fitz."):
            raise ImportError("No module named 'mupdf'")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    class _FakePage:
        def extract_text(self):
            return "降级提取文本"

    class _FakePdf:
        def __init__(self):
            self.pages = [_FakePage()]

        def __enter__(self):
            return self

        def __exit__(self, *exc):
            return False

    fake_pdfplumber = types.SimpleNamespace(open=lambda _path: _FakePdf())
    monkeypatch.setitem(sys.modules, "pdfplumber", fake_pdfplumber)

    pages = asyncio.run(document_processor._extract_pdf("dummy.pdf"))
    assert pages == [PageSegment("降级提取文本", 1)]


def test_clean_text_strips_garbage():
    """_clean_text 应剔除中文 PDF 文本层常见的乱码码位，避免引用来源出现「�」。"""
    # C1 控制符（U+0085 下一行 / U+0096 等）应被剔除，保留可见文字与换行
    assert _clean_text("正文\u0085内容\n第二行") == "正文内容\n第二行"
    # Unicode 替换符 U+FFFD
    assert _clean_text("知识\uFFFD库内容") == "知识库内容"
    # 非字符码位 U+FFFE / U+FDEF
    assert _clean_text("a\ufffeb") == "ab"
    assert _clean_text("x\ufdefy") == "xy"
    # 孤立代理对（单个代理单元）应剔除，不污染后续字符
    assert _clean_text("前\ud83d后") == "前后"
    # 合法 emoji（完整代理对）应保留
    assert _clean_text("点赞\U0001f44d") == "点赞\U0001f44d"
    # BOM 与制表/换行保留
    assert _clean_text("\ufeff\t\t标题\r\n正文") == "\t\t标题\r\n正文"
    # 空输入
    assert _clean_text("") == ""


def test_chunk_text_parent_child_layers():
    """M5-5：启用父子分块时，chunk_text 产出 parent 块与子块两层，且子块带 parent_id 归属。"""
    # 用较长文本确保能切出多个父块/子块
    text = ("熊答是企业知识问答助手，支持 RAG 检索与智能问答，具备多租户与成员管理能力。" * 30)
    pages = [PageSegment(text, 1)]
    with pytest.MonkeyPatch().context() as mp:
        mp.setattr(dp_module.settings, "enable_parent_child", True)
        chunks = asyncio.run(document_processor.chunk_text(pages))

    parents = [c for c in chunks if c.metadata.get("chunk_type") == "parent"]
    children = [c for c in chunks if c.metadata.get("chunk_type") != "parent"]
    assert len(parents) >= 1, "应至少产出一个父块"
    assert len(children) >= 1, "应至少产出子块"
    # 每个子块都应有 parent_id 归属
    assert all(c.metadata.get("parent_id") is not None for c in children)
    # 父块不应与子块内容完全重复（父块更大）
    assert all(len(p.content) >= len(c.content) for p in parents for c in children)


def test_chunk_text_parent_child_disabled_fallback():
    """M5-5：关闭父子分块开关时，退化为原单层分块（无 parent 块，子块无 parent_id）。"""
    text = ("熊答是企业知识问答助手，支持 RAG 检索与智能问答。" * 10)
    pages = [PageSegment(text, 1)]
    with pytest.MonkeyPatch().context() as mp:
        mp.setattr(dp_module.settings, "enable_parent_child", False)
        chunks = asyncio.run(document_processor.chunk_text(pages))

    assert all(c.metadata.get("chunk_type") != "parent" for c in chunks)
    assert all("parent_id" not in c.metadata for c in chunks)


def test_embed_chunks_skips_parent():
    """M5-5：embed_chunks 跳过父块向量化，仅向量化非 parent 块。"""
    from services.document_processor import DocumentChunk

    chunks = [
        DocumentChunk(
            content="父块完整上下文",
            metadata={"chunk_index": 0, "chunk_type": "parent", "is_parent": True, "parent_id": "p0"},
        ),
        DocumentChunk(
            content="子块内容",
            metadata={"chunk_index": 1, "chunk_type": "child", "parent_id": "p0"},
        ),
    ]

    async def fake_embed_batch(texts, cfg=None):
        return [[0.1, 0.2, 0.3] for _ in texts]

    with pytest.MonkeyPatch().context() as mp:
        mp.setattr(dp_module.embedding_service, "embed_batch", fake_embed_batch)
        out = asyncio.run(document_processor.embed_chunks(chunks))
    # 父块 embedding 保持 None（不进向量索引）
    assert out[0].embedding is None
    # 子块被向量化
    assert out[1].embedding is not None


def test_generate_extended_augment_four_types():
    """M5-7：_generate_extended_augment 生成 summary/question/wiki/entity 四类块，entity 带三元组。"""
    from services.document_processor import DocumentChunk

    chunks = [
        DocumentChunk(
            content="熊答是企业知识问答助手，支持 RAG 检索与智能问答。",
            metadata={"chunk_index": 0, "page": 1},
        )
    ]

    async def fake_complete(messages, model=None, cfg=None, client=None):
        system = messages[0]["content"]
        if "摘要" in system:
            return "熊答是企业知识问答助手。"
        if "Auto-Wiki" in system:
            return '["条目A：熊答定义", "条目B：RAG 检索"]'
        if "实体关系" in system:
            return '[{"subject":"熊答","predicate":"属于","object":"企业知识库"}]'
        if "检索式问题" in system:
            return '["什么是熊答？", "熊答支持什么能力？"]'
        return "{}"

    async def fake_embed_batch(texts, cfg=None, client=None):
        return [[0.1, 0.2, 0.3] for _ in texts]

    with pytest.MonkeyPatch().context() as mp:
        mp.setattr(dp_module.llm_service, "complete", fake_complete)
        mp.setattr(dp_module.embedding_service, "embed_batch", fake_embed_batch)
        out = asyncio.run(document_processor._generate_extended_augment(chunks, None))

    types = {d["metadata"]["chunk_type"] for d in out}
    assert types == {"summary", "question", "wiki", "entity"}, types
    # 每块都被向量化
    assert all(d["embedding"] is not None for d in out)
    # entity 块携带结构化三元组（存 metadata，供图检索/可视化）
    entity = [d for d in out if d["metadata"]["chunk_type"] == "entity"][0]
    assert entity["metadata"]["triple"] == {
        "subject": "熊答",
        "predicate": "属于",
        "object": "企业知识库",
    }
    assert "企业知识库" in entity["content"]


def test_generate_extended_augment_disabled():
    """M5-7：总开关 enable_augment_extensions 关闭时返回空列表。"""
    from services.document_processor import DocumentChunk

    chunks = [DocumentChunk(content="内容", metadata={"chunk_index": 0, "page": 1})]

    async def fake_embed_batch(texts, cfg=None, client=None):
        return [[0.1] for _ in texts]

    with pytest.MonkeyPatch().context() as mp:
        mp.setattr(dp_module.settings, "enable_augment_extensions", False)
        mp.setattr(dp_module.embedding_service, "embed_batch", fake_embed_batch)
        out = asyncio.run(document_processor._generate_extended_augment(chunks, None))
    assert out == []


def test_generate_extended_augment_model_config_error():
    """M5-7：LLM 配置错误时整体跳过扩展增强（返回空列表）。"""
    from services.document_processor import DocumentChunk
    from services.model_config import ModelConfigError

    chunks = [DocumentChunk(content="内容", metadata={"chunk_index": 0, "page": 1})]

    async def fake_complete_err(messages, model=None, cfg=None, client=None):
        raise ModelConfigError("bad key")

    async def fake_embed_batch(texts, cfg=None, client=None):
        return [[0.1] for _ in texts]

    with pytest.MonkeyPatch().context() as mp:
        mp.setattr(dp_module.llm_service, "complete", fake_complete_err)
        mp.setattr(dp_module.embedding_service, "embed_batch", fake_embed_batch)
        out = asyncio.run(document_processor._generate_extended_augment(chunks, None))
    assert out == []


def test_extract_images_dispatch():
    """M5-6：extract_images 对 txt/未知类型返回 []，对 pdf/docx 委派到对应抽取器。"""
    with pytest.MonkeyPatch().context() as mp:
        mp.setattr(
            document_processor,
            "_extract_pdf_images",
            lambda fp: [{"bytes": b"p", "page": 1, "ext": "png"}],
        )
        mp.setattr(
            document_processor,
            "_extract_docx_images",
            lambda fp: [{"bytes": b"d", "page": 0, "ext": "png"}],
        )
        assert asyncio.run(document_processor.extract_images("x.pdf", "pdf")) == [
            {"bytes": b"p", "page": 1, "ext": "png"}
        ]
        assert asyncio.run(document_processor.extract_images("x.docx", "docx")) == [
            {"bytes": b"d", "page": 0, "ext": "png"}
        ]
        # 非图片类文档（txt/md）直接返回空，不触发抽取
        assert asyncio.run(document_processor.extract_images("x.txt", "txt")) == []
        assert asyncio.run(document_processor.extract_images("x.md", "md")) == []


def test_extract_docx_images_real_zip():
    """M5-6：_extract_docx_images 从 zip(word/media/) 抽取图片，跳过过小的图片。"""
    import zipfile

    small = b"\x89PNG\r\n\x1a\n" + b"\x00" * 10  # 小于 min bytes，应被过滤
    big = b"\x89PNG\r\n\x1a\n" + b"\x00" * 5000  # 大于 min bytes，保留

    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("word/document.xml", "<w:document/>")
        z.writestr("word/media/image1.png", small)
        z.writestr("word/media/image2.png", big)
    try:
        imgs = document_processor._extract_docx_images(path)
        # 仅大图被保留（小图被 min bytes 过滤）
        assert len(imgs) == 1
        assert imgs[0]["ext"] == "png"
        assert imgs[0]["page"] == 0  # DOCX 无真实页码
    finally:
        os.unlink(path)


def test_generate_multimodal_augment_ocr_caption():
    """M5-6：_generate_multimodal_augment 对图片产出 ocr + image_caption 块并向量化。"""
    from services.vector_store import AUGMENT_CHUNK_TYPES

    # 真实临时文件，使 _generate_multimodal_augment 的 os.path.exists 守卫通过
    fd, img_pdf = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    task = {"file_path": img_pdf, "file_type": "pdf"}
    try:

        async def fake_extract_images(fp, ft):
            return [{"bytes": b"imgbytes1", "page": 2, "ext": "png"}]

        async def fake_complete(messages, model=None, cfg=None, client=None):
            # 多模态消息 content 为 parts 列表（含 image_url），此处直接返回 OCR 结果
            return '{"ocr": "图中显示登录流程图", "caption": "一张描述系统登录流程的示意图"}'

        async def fake_embed_batch(texts, cfg=None, client=None):
            return [[0.1, 0.2, 0.3] for _ in texts]

        with pytest.MonkeyPatch().context() as mp:
            mp.setattr(dp_module.settings, "enable_multimodal", True)
            mp.setattr(document_processor, "extract_images", fake_extract_images)
            mp.setattr(dp_module.llm_service, "complete", fake_complete)
            mp.setattr(dp_module.embedding_service, "embed_batch", fake_embed_batch)
            out = asyncio.run(document_processor._generate_multimodal_augment(task, None))
    finally:
        os.unlink(img_pdf)

    types = {d["metadata"]["chunk_type"] for d in out}
    assert types == {"ocr", "image_caption"}, types
    # 每块都被向量化
    assert all(d["embedding"] is not None for d in out)
    ocr = [d for d in out if d["metadata"]["chunk_type"] == "ocr"][0]
    assert ocr["content"] == "图中显示登录流程图"
    assert ocr["metadata"]["page"] == 2
    # 两块均被排除为引用来源（chunk_type 在 AUGMENT_CHUNK_TYPES）
    assert "ocr" in AUGMENT_CHUNK_TYPES and "image_caption" in AUGMENT_CHUNK_TYPES


def test_generate_multimodal_augment_no_images():
    """M5-6：文档无图片时返回空列表（不影响原有增强）。"""
    task = {"file_path": "x.txt", "file_type": "txt"}  # txt 不抽图

    with pytest.MonkeyPatch().context() as mp:
        mp.setattr(dp_module.settings, "enable_multimodal", True)
        out = asyncio.run(document_processor._generate_multimodal_augment(task, None))
    assert out == []


def test_generate_multimodal_augment_disabled():
    """M5-6：总开关 enable_multimodal 关闭时返回空列表。"""
    task = {"file_path": "x.pdf", "file_type": "pdf"}

    with pytest.MonkeyPatch().context() as mp:
        mp.setattr(dp_module.settings, "enable_multimodal", False)
        out = asyncio.run(document_processor._generate_multimodal_augment(task, None))
    assert out == []


def test_generate_multimodal_augment_model_config_error():
    """M5-6：VLM 配置错误时整体跳过多模态增强（返回空列表）。"""
    from services.model_config import ModelConfigError

    task = {"file_path": "x.pdf", "file_type": "pdf"}

    async def fake_extract_images(fp, ft):
        return [{"bytes": b"img", "page": 1, "ext": "png"}]

    async def fake_complete(messages, model=None, cfg=None, client=None):
        raise ModelConfigError("bad key")

    with pytest.MonkeyPatch().context() as mp:
        mp.setattr(dp_module.settings, "enable_multimodal", True)
        mp.setattr(document_processor, "extract_images", fake_extract_images)
        mp.setattr(dp_module.llm_service, "complete", fake_complete)
        out = asyncio.run(document_processor._generate_multimodal_augment(task, None))
    assert out == []






