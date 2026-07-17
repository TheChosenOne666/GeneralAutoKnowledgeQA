"""文档处理服务 — 提取 → 分块 → 向量化 → 存储。

使用 LangChain 的文档加载器和文本分割器。
M2-3：完整链路已接入 Embedding 与 VectorStore（M1-5 阶段仅提取+分块）。

页码策略：
- PDF：PyMuPDF 逐页提取，使用真实页码。
- DOCX / TXT / MD：无真实页码，按 CHARS_PER_PAGE 估算（仅用于引用来源展示）。
"""

import asyncio
import json
import os
import re
from dataclasses import dataclass

from loguru import logger

from core.config import settings
from services import augment_queue
from services import process_queue
from services import vector_store as vector_store_module
from services.embedding import embedding_service
from services.llm import llm_service
from services.model_config import ModelConfig, ModelConfigError
from services.status_callback import notify_document_status

# DOCX/TXT/MD 无真实页码，按字符数估算每页（仅用于引用来源展示，非精确）
CHARS_PER_PAGE = 1500


@dataclass
class DocumentChunk:
    """文档分块。"""

    content: str
    metadata: dict
    embedding: list[float] | None = None


@dataclass
class PageSegment:
    """按页分段的文本。"""

    text: str
    page: int


# 需剔除的「不可见 / 非字符」Unicode 区间：
# - C0 控制符（保留 \n \r \t）：U+0000–U+0008、U+000B、U+000C、U+000E–U+001F
# - DEL 与 C1 控制符：U+007F–U+009F（中文 PDF 文本层常见，前端渲染成「�」/方框）
# - 非字符码位：U+FDD0–U+FDEF、U+FFFE、U+FFFF
# 这些字符不应进入分块 / 引用来源；仅保留 \n \r \t 与可见字符。
_GARBAGE_RE = re.compile(r"[\u0000-\u0008\u000b\u000c\u000e-\u001f\u007f-\u009f\ufdd0-\ufdef\ufffe\uffff]")


def _clean_text(text: str) -> str:
    """清洗提取文本：去 BOM、去控制/非字符码位、去替换符、去孤立代理对。

    中文 PDF 文本层常含 C1 控制符或无法映射的字形（PyMuPDF 以 U+FFFD 占位），
    若不过滤会进入分块与引用来源，前端渲染成「�」乱码。仅保留 \\n \\r \\t 与可见字符。
    """
    if not text:
        return ""
    text = text.replace("\ufeff", "")          # BOM
    text = _GARBAGE_RE.sub("", text)            # C0/C1 控制符 + 非字符码位
    text = text.replace("\ufffd", "")           # Unicode 替换符（解析乱码占位）
    # 剔除孤立代理对（surrogate），避免存储 / 截断产生 U+FFFD 乱码
    text = "".join(ch for ch in text if not (0xD800 <= ord(ch) <= 0xDFFF))
    return text


class DocumentProcessor:
    """文档处理流水线。"""

    @staticmethod
    def _read_text_file(file_path: str) -> str:
        """读取文本文件，依次尝试 utf-8 / gbk，避免中文 Windows 文件乱码。"""
        with open(file_path, "rb") as f:
            data = f.read()
        for enc in ("utf-8", "gbk", "latin-1"):
            try:
                return data.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue
        return data.decode("utf-8", errors="replace")

    @staticmethod
    def _estimate_pages(text: str) -> list[PageSegment]:
        """按 CHARS_PER_PAGE 估算页码（DOCX/TXT/MD 无真实页码）。"""
        text = text or ""
        if not text.strip():
            return [PageSegment(text, 1)]
        segs: list[PageSegment] = []
        for i in range(0, len(text), CHARS_PER_PAGE):
            segs.append(PageSegment(text[i : i + CHARS_PER_PAGE], len(segs) + 1))
        return segs

    async def extract_pages(self, file_path: str, file_type: str) -> list[PageSegment]:
        """提取文本并按页分段，返回 [(text, page_no), ...]。

        - PDF: PyMuPDF 逐页真实页码
        - DOCX: python-docx（正文 + 表格），无真实页码时按字符数估算
        - MD/TXT: 按字符数估算页码
        """
        ft = file_type.lower()
        if ft == "pdf":
            return await self._extract_pdf(file_path)
        if ft == "docx":
            from docx import Document

            doc = Document(file_path)
            parts = [p.text for p in doc.paragraphs]
            # 表格内容也纳入提取，避免正文缺失
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    parts.append(" | ".join(cells))
            full = _clean_text("\n".join(parts))
            return self._estimate_pages(full)
        # md / txt / 其他：直接读取（编码容错）
        raw = self._read_text_file(file_path)
        return self._estimate_pages(_clean_text(raw))

    async def _extract_pdf(self, file_path: str) -> list[PageSegment]:
        """提取 PDF 文本并逐页分段（真实页码）。

        优先 PyMuPDF（快、逐页真实页码）；当环境缺少 C 扩展（如部分 Windows 上
        ``_mupdf.pyd`` 因 VC++ 运行库 / DLL 加载失败）导致 ``import fitz`` 不可用时，
        自动降级到纯 Python 的 pdfplumber，保证 PDF 在该环境下仍可被处理。
        """
        try:
            import fitz
        except ImportError as e:
            logger.warning(f"[PDF] PyMuPDF 不可用（{e}），降级到 pdfplumber")
            return await self._extract_pdf_pdfplumber(file_path)
        doc = fitz.open(file_path)
        try:
            segs = [
                PageSegment(_clean_text(page.get_text()), i + 1)
                for i, page in enumerate(doc)
            ]
        finally:
            doc.close()
        return segs

    async def _extract_pdf_pdfplumber(self, file_path: str) -> list[PageSegment]:
        """pdfplumber 降级提取（纯 Python，无 C 扩展依赖），重 IO/CPU 放到线程池。"""
        import pdfplumber

        def _run() -> list[PageSegment]:
            segs: list[PageSegment] = []
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    segs.append(PageSegment(_clean_text(text), i + 1))
            return segs

        return await asyncio.to_thread(_run)

    async def chunk_text(self, pages: list[PageSegment]) -> list[DocumentChunk]:
        """文本分块 — 使用 LangChain RecursiveCharacterTextSplitter，保留每块的页码。

        方案C（大纲感知）：除常规内容块外，额外扫描每页文本行，识别章节/小节
        标题并单独生成 ``chunk_type='outline'`` 的标题块（记录 ``section_title``）。
        这些大纲块供「知识架构/大纲/目录」类问题在检索时优先召回，让 LLM 能基于
        文档真实章节主题归纳出知识框架（见 :mod:`services.rag` 的大纲意图路由）。
        """
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
        )
        chunks: list[DocumentChunk] = []
        global_index = 0
        for seg in pages:
            if not seg.text.strip():
                continue
            for c in splitter.split_text(seg.text):
                # 防御性再清洗：即便上游提取漏清，落库 chunk 也一定干净（避免 U+FFFD 进入引用来源）
                chunks.append(
                    DocumentChunk(
                        content=_clean_text(c),
                        metadata={"chunk_index": global_index, "page": seg.page},
                    )
                )
                global_index += 1
            # 大纲感知：扫描该页标题行，生成 outline 块（不破坏内容块，仅补充）
            for title in self._extract_outline_titles(seg.text):
                chunks.append(
                    DocumentChunk(
                        content=_clean_text(title),
                        metadata={
                            "chunk_index": global_index,
                            "page": seg.page,
                            "chunk_type": "outline",
                            "section_title": title,
                        },
                    )
                )
                global_index += 1
        return chunks

    @staticmethod
    def _extract_outline_titles(text: str) -> list[str]:
        """从文本中识别章节/小节标题行，返回去重后的标题列表。

        启发式（针对中文文档 / PDF 文本层）：短行（<=50 字）+ 标题特征
        （数字/中文编号、章节词、架构/大纲/目录词、追问式短句）。用于生成
        ``outline`` 分块，供大纲意图检索优先召回。
        """
        titles: list[str] = []
        seen: set[str] = set()
        for line in text.split("\n"):
            t = line.strip()
            if not t or len(t) > 50:
                continue
            if DocumentProcessor._is_title_line(t) and t not in seen:
                seen.add(t)
                titles.append(t)
        return titles

    @staticmethod
    def _is_title_line(t: str) -> bool:
        """判断单行文本是否为章节/小节标题（启发式）。

        仅保留「编号章节」（1. / 1.2 / (1) / 一、 / 第X章）与明显的大纲词
        （架构/大纲/目录/知识点/考点/体系），不把具体考点追问句与冒号短句
        当作标题，避免 outline 块过多稀释常规检索。
        """
        # 数字/字母编号：1. 1.2 (1) 一、二、第X章/节/篇
        if re.match(
            r"^(\d+(\.\d+)*|[（(]\d+[）)]|[一二三四五六七八九十]+、|"
            r"第[一二三四五六七八九十\d]+[章节篇])",
            t,
        ):
            return True
        # 含架构/大纲/目录/知识点/考点/框架等关键词的短句
        if any(kw in t for kw in ("架构", "大纲", "目录", "知识点", "考点", "知识框架", "体系")):
            return True
        return False

    async def embed_chunks(
        self, chunks: list[DocumentChunk], cfg: ModelConfig | None = None
    ) -> list[DocumentChunk]:
        """向量化分块。"""
        texts = [c.content for c in chunks]
        vectors = await embedding_service.embed_batch(texts, cfg)
        for chunk, vec in zip(chunks, vectors):
            chunk.embedding = vec
        return chunks

    async def process(
        self,
        file_path: str,
        file_type: str,
        kb_id: str,
        doc_id: str,
        tenant_id: str,
        cfg: ModelConfig | None = None,
        ai_config_dict: dict | None = None,
    ) -> dict:
        """完整文档处理流水线：提取 → 分块 → 检索(向量化+入库) → 优化(问答增强·后台异步)。

        对齐 业界 finalizing（异步增强） 语义：向量化完成并入库原始块后，文档**立即可被检索**，
        随后进入 ``optimizing``（增强问答对中）并由后台任务异步跑增强，HTTP 立即返回
        ``optimizing``，不再阻塞等待整篇增强完成（避免界面长时间卡在「优化中」）。

        阶段状态通过 :func:`notify_document_status` 回调 Java 实时透出：
        解析(由 Java 置) → 检索(retrieving) → 优化(optimizing，已可检索) → 就绪(后台回调)。

        Returns:
            {"status": "optimizing", "chunk_count": int, "content": str}
                — status 恒为 optimizing（向量已入库可检索，增强后台进行中）；
                  content 为提取全文，供前端查看。
        """
        # 1. 提取文本（按页）
        logger.info(f"[文档诊断] 开始提取文本 doc_id={doc_id} file_type={file_type} file_path={file_path}")
        pages = await self.extract_pages(file_path, file_type)
        full_text = _clean_text("\n".join(p.text for p in pages))
        logger.info(f"[文档诊断] 提取完成 doc_id={doc_id} 文本长度={len(full_text)} 页数={len(pages)}")

        # 2. 分块（保留页码）
        chunks = await self.chunk_text(pages)
        logger.info(f"[文档诊断] 分块完成 doc_id={doc_id} 块数={len(chunks)}")

        # 3. 检索阶段：向量化 + 构建原始块（使文档可被检索）
        await notify_document_status(doc_id, "retrieving")
        # 取消检查点（嵌入前）：用户中途取消则不消耗 Embedding 调用、不入库
        if await augment_queue.is_cancelled(doc_id):
            logger.info(f"[文档诊断] 处理中被取消（嵌入前），放弃 doc_id={doc_id}")
            await notify_document_status(doc_id, "cancelled")
            return {"status": "cancelled", "chunk_count": len(chunks), "content": full_text}
        chunks = await self.embed_chunks(chunks, cfg)
        source = os.path.basename(file_path)
        chunk_dicts = []
        for c in chunks:
            meta = dict(c.metadata)
            meta.update(
                {
                    "doc_id": doc_id,
                    "kb_id": kb_id,
                    "tenant_id": tenant_id,
                    "source": source,
                }
            )
            chunk_dicts.append(
                {
                    "content": c.content,
                    "metadata": meta,
                    "embedding": c.embedding,
                }
            )

        # 4. 原始块先入库 —— 向量一旦入库，文档即可被检索（对齐 业界 finalizing（异步增强）=queryable）
        await vector_store_module.vector_store_service.store_chunks(
            chunk_dicts, kb_id, doc_id, tenant_id
        )
        logger.info(f"[文档诊断] 原始块入库完成 doc_id={doc_id} 块数={len(chunk_dicts)}（已可检索）")

        # 5. 优化阶段：取消检查点（入库后、notify optimizing 之前）
        # 用户中途取消 → 清理已写向量 + 回调 cancelled，避免广播「已可检索」假提示
        # （向量已入库可检索，必须清理避免残留「已取消却可检索」的孤立向量）
        if await augment_queue.is_cancelled(doc_id):
            logger.info(f"[文档诊断] 处理中被取消（入库后），清理向量 doc_id={doc_id}")
            await vector_store_module.vector_store_service.delete_by_doc(doc_id)
            await notify_document_status(doc_id, "cancelled")
            return {"status": "cancelled", "chunk_count": len(chunks), "content": full_text}
        # 进入 optimizing（已可检索），将增强任务入持久化队列，HTTP 立即返回
        # M5-1：全文经状态回调回填 Java（替代旧同步返回），供前端「查看内容」弹窗
        await notify_document_status(doc_id, "optimizing", content=full_text, chunk_count=len(chunks))
        if settings.enable_qa_augment:
            # 任务入 Redis 队列（持久化，worker 常驻消费，服务重启可由 sweep 恢复）
            await augment_queue.enqueue(
                {
                    "doc_id": doc_id,
                    "kb_id": kb_id,
                    "tenant_id": tenant_id,
                    "file_path": file_path,
                    "file_type": file_type,
                    "ai_config": ai_config_dict,
                }
            )
        else:
            # 未开启增强：直接回调 ready（无需再入库，原始块已在库）
            await notify_document_status(doc_id, "ready")

        return {"status": "optimizing", "chunk_count": len(chunks), "content": full_text}

    async def _run_augment_task(self, task: dict) -> None:
        """持久化队列消费：执行单文档问答增强并回调 ready（对齐 业界 finalizing（异步增强） 队列任务）。

        任务来自 Redis 队列（可跨进程重启恢复）：
        - 取消检查：文档已删（cancelled 集命中 / 原始块已不存在）则跳过，不回调 ready；
        - 原始块重建：从向量库读回（不依赖上传文件是否在），重 embed 后连同增强块全量 store；
        - 看门狗：总超时 + 异常兜底，保证最终推进 ready（向量化早已完成、可检索）。
        """
        doc_id = task.get("doc_id")
        kb_id = task.get("kb_id")
        tenant_id = task.get("tenant_id")
        # 防御：字段缺失的畸形任务（如历史上测试误写入 Redis 的脏数据）直接跳过并 ack 移除，
        # 避免 worker 每轮重复崩溃刷屏。
        if not (doc_id and kb_id and tenant_id):
            logger.warning(f"[优化] 任务字段缺失，跳过并移除 doc_id={doc_id}: {task}")
            return
        ai_config_dict = task.get("ai_config")
        cfg = ModelConfig.from_dict(ai_config_dict) if ai_config_dict else None
        skip_ready = False
        try:
            # 取消检查（文档被删除时由 Java 调 Python 标记）
            if await augment_queue.is_cancelled(doc_id):
                logger.info(f"[优化] 任务已取消，跳过 doc_id={doc_id}")
                skip_ready = True
                return
            # 重建原始块（从向量库读回，避免依赖上传文件是否仍在）
            originals = await vector_store_module.vector_store_service.get_original_chunks(doc_id)
            if not originals:
                logger.warning(f"[优化] 文档向量已不存在（可能已删除），放弃增强 doc_id={doc_id}")
                skip_ready = True
                return

            source = originals[0].get("source", "")
            chunk_dicts: list[dict] = []
            chunks: list[DocumentChunk] = []
            for o in originals:
                meta = {
                    "doc_id": doc_id,
                    "kb_id": kb_id,
                    "tenant_id": tenant_id,
                    "source": o.get("source", source),
                    "chunk_index": o.get("chunk_index", 0),
                    "page": o.get("page", 0),
                }
                chunk_dicts.append({"content": o["content"], "metadata": meta, "embedding": None})
                chunks.append(
                    DocumentChunk(
                        content=o["content"],
                        metadata={"chunk_index": o.get("chunk_index", 0), "page": o.get("page", 0)},
                    )
                )
            # 重 embed 原始块（命中 L2 缓存，近乎零成本），供最终全量 store
            texts = [c["content"] for c in chunk_dicts]
            vecs = await embedding_service.embed_batch(texts, cfg)
            for c, vec in zip(chunk_dicts, vecs):
                c["embedding"] = vec
            for c, vec in zip(chunks, vecs):
                c.embedding = vec

            augmented = await asyncio.wait_for(
                self._generate_qa_augment(chunks, cfg),
                timeout=settings.qa_background_timeout,
            )
            if augmented:
                await vector_store_module.vector_store_service.store_chunks(
                    chunk_dicts + augmented, kb_id, doc_id, tenant_id
                )
                logger.info(f"[优化] 增强入库完成 doc_id={doc_id} 增强块={len(augmented)}")
            else:
                logger.info(f"[优化] 无增强块（LLM 不可用/配置错误），文档维持可检索 doc_id={doc_id}")
        except asyncio.TimeoutError:
            logger.warning(f"[优化] 问答增强总超时，放弃剩余块，文档维持可检索 doc_id={doc_id}")
        except ModelConfigError as e:
            logger.warning(f"[优化] LLM/Embedding 配置错误，跳过增强 doc_id={doc_id}: {e}")
        except Exception as e:
            logger.warning(f"[优化] 后台增强异常（文档已可检索，忽略）doc_id={doc_id}: {e}")
        finally:
            # 始终从 processing 移除（无论成功/失败/取消/删除）
            try:
                await augment_queue.ack(task)
            except Exception as e:
                logger.warning(f"[优化] ack 任务失败 doc_id={doc_id}: {e}")
            # 仅非取消/删除时回调 ready（向量化早已完成，可检索）
            if not skip_ready:
                try:
                    await notify_document_status(doc_id, "ready")
                except Exception as e:
                    logger.warning(f"[优化] 回调 ready 失败 doc_id={doc_id}: {e}")

    async def run_augment_worker(self, poll_interval: float = 2.0) -> None:
        """常驻消费增强队列（在 FastAPI lifespan 中以 asyncio 任务启动）。

        单 worker 串行消费（增强为 LLM 密集，并发由 ``_generate_qa_augment`` 内部限并发控制）。
        队列持久化于 Redis，进程重启后由 lifespan 的 ``sweep_stale`` 恢复卡死任务。
        """
        logger.info("[优化] 增强队列 worker 启动")
        while True:
            try:
                task = await augment_queue.dequeue()
                if task is None:
                    await asyncio.sleep(poll_interval)
                    continue
                await self._run_augment_task(task)
            except Exception as e:
                logger.warning(f"[优化] worker 消费异常（继续下一轮）: {e}")
                await asyncio.sleep(poll_interval)

    async def run_process_worker(self, poll_interval: float = 2.0) -> None:
        """常驻消费主流程队列（M5-1，在 FastAPI lifespan 中以 asyncio 任务启动）。

        消费 ``process_queue`` 中的任务，执行 ``process()``（提取 → 分块 → 向量化 → 入库 →
        入增强队）。状态完全经 ``notify_document_status`` 回调 Java 推进（retrieving /
        optimizing / ready / failed），HTTP 上传仅 enqueue 不阻塞。

        队列持久化于 Redis，进程重启后由 lifespan 的 ``process_queue.sweep_stale`` 恢复卡死任务。
        """
        logger.info("[主流程] 文档处理队列 worker 启动")
        while True:
            try:
                task = await process_queue.dequeue()
                if task is None:
                    await asyncio.sleep(poll_interval)
                    continue
                await self._run_process_task(task)
            except Exception as e:
                logger.warning(f"[主流程] worker 消费异常（继续下一轮）: {e}")
                await asyncio.sleep(poll_interval)

    async def _run_process_task(self, task: dict) -> None:
        """执行单个主流程任务（M5-1），失败回调 Java failed。"""
        doc_id = task.get("doc_id")
        try:
            if await process_queue.is_cancelled(doc_id):
                logger.info(f"[主流程] 任务已取消，跳过 doc_id={doc_id}")
                return
            # 从用户级 ai_config 构建 cfg（与原同步 handler 一致，复用用户/租户模型配置），
            # 而非走 env 兜底；ai_config_dict 一并透传供 process() 内部使用。
            cfg = ModelConfig.from_dict(task.get("ai_config"))
            # 复用现有 process()：提取→分块→向量化→入库→入增强队，
            # 内部已含取消检查与 retrieving/optimizing 状态回调；返回值为历史兼容，忽略。
            await self.process(
                file_path=task["file_path"],
                file_type=task["file_type"],
                kb_id=task["kb_id"],
                doc_id=doc_id,
                tenant_id=task["tenant_id"],
                cfg=cfg,
                ai_config_dict=task.get("ai_config"),
            )
        except ModelConfigError as e:
            logger.warning(f"[主流程] 模型配置错误 doc_id={doc_id}: {e}")
            await notify_document_status(
                doc_id, "failed", error_msg=str(e), model_config_error=True
            )
        except Exception as e:
            logger.exception(f"[主流程] 文档处理失败 doc_id={doc_id}: {e}")
            await notify_document_status(doc_id, "failed", error_msg=str(e)[:500])
        finally:
            try:
                await process_queue.ack(task)
            except Exception as e:
                logger.warning(f"[主流程] ack 失败 doc_id={doc_id}: {e}")

    async def _generate_qa_augment(
        self, chunks: list[DocumentChunk], cfg: ModelConfig | None = None, client=None
    ) -> list[dict]:
        """基于原始分块并发生成问答对，批量向量化后作为增强块（retrieval augmentation）。

        对标业界成熟方案：每批并发生成（限并发度）+ 批量 embedding（一次调用），而非单线程串行。
        best-effort：LLM 不可用 / 配置错误 / 单块超时，跳过剩余增强（文档仍就绪可用），
        不阻塞主流程。
        """
        # 跳过 outline 块（章节标题），仅对内容块生成问答增强，避免短标题污染增强集
        selected = [
            c for c in chunks if c.metadata.get("chunk_type") != "outline"
        ][: settings.qa_max_pairs]
        augmented: list[dict] = []
        results: list[tuple[int, dict | None, bool]] = [(i, None, False) for i in range(len(selected))]
        semaphore = asyncio.Semaphore(settings.qa_concurrency)

        async def gen_one(i: int, c: DocumentChunk) -> None:
            async with semaphore:
                try:
                    qa = await asyncio.wait_for(
                        self._gen_one_qa(c.content, cfg, client),
                        timeout=settings.qa_per_qa_timeout,
                    )
                    results[i] = (i, qa, False)
                except ModelConfigError:
                    # 标记配置错误，外层整体跳过增强
                    results[i] = (i, None, True)
                except Exception as e:
                    logger.warning(f"[优化] 生成单个问答对失败，跳过该块 doc_id块{i}: {e}")
                    results[i] = (i, None, False)

        await asyncio.gather(*(gen_one(i, c) for i, c in enumerate(selected)))

        # 任一配置错误则整体跳过增强（与 M3-3「模型配置错误」语义一致）
        if any(flag for _, _, flag in results):
            logger.warning("[优化] 检测到模型配置错误，跳过整篇问答增强")
            return []

        valid = [(i, qa) for i, qa, _ in results if qa]
        if not valid:
            return []

        # 收集所有问题，一次性批量 embedding（减少 HTTP 往返，命中 L2 缓存）
        questions = [qa["question"] for _, qa in valid]
        try:
            vecs = await embedding_service.embed_batch(questions, cfg, client)
        except ModelConfigError as e:
            logger.warning(f"[优化] Embedding 配置错误，跳过增强: {e}")
            return []

        for (idx, qa), vec in zip(valid, vecs):
            c = selected[idx]
            augmented.append(
                {
                    "content": f"Q: {qa['question']}\nA: {qa['answer']}",
                    "metadata": {
                        "chunk_index": c.metadata.get("chunk_index", 0),
                        "page": c.metadata.get("page", 0),
                        "chunk_type": "qa",
                    },
                    "embedding": vec,
                }
            )
        return augmented

    @staticmethod
    async def _gen_one_qa(text: str, cfg: ModelConfig | None, client=None) -> dict | None:
        """用 LLM 从文本生成一对问答，返回 ``{"question", "answer"}`` 或 ``None``。"""
        messages = [
            {
                "role": "system",
                "content": (
                    "你是知识库检索优化助手。阅读给定文本，生成一个最能用于检索该知识点的用户提问"
                    "及其准确、简洁的答案。只输出一个 JSON 对象："
                    '{"question": "...", "answer": "..."}，不要包含任何额外文字或代码块标记。'
                ),
            },
            {"role": "user", "content": text[:2000]},
        ]
        try:
            raw = await llm_service.complete(messages, cfg=cfg, client=client)
        except ModelConfigError:
            raise
        except Exception as e:
            logger.warning(f"[优化] LLM 调用异常: {e}")
            return None
        return _parse_qa_json(raw)


def _parse_qa_json(raw: str) -> dict | None:
    """从 LLM 输出中容错解析问答 JSON，返回 ``{"question", "answer"}`` 或 ``None``。"""
    if not raw:
        return None
    text = raw.strip()
    # 去除可能的 ```json ... ``` 包裹
    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("json"):
            text = text[4:]
    # 1) 尝试整段直接解析
    try:
        data = json.loads(text)
        if isinstance(data, dict) and data.get("question") and data.get("answer"):
            return {"question": str(data["question"]).strip(), "answer": str(data["answer"]).strip()}
    except (json.JSONDecodeError, TypeError):
        pass
    # 2) 兜底：截取第一个 {...} 片段解析
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            data = json.loads(text[start : end + 1])
            if isinstance(data, dict) and data.get("question") and data.get("answer"):
                return {
                    "question": str(data["question"]).strip(),
                    "answer": str(data["answer"]).strip(),
                }
        except (json.JSONDecodeError, TypeError):
            pass
    return None


document_processor = DocumentProcessor()
